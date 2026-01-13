"""DOCX document processor for PII redaction."""

import os
import copy
from typing import List, Tuple, Optional, Callable
from dataclasses import dataclass

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

from ..detectors.pii_detector import PIIDetector, DetectionResult
from ..detectors.patterns import PIIMatch, PIIType
from ..storage.mapping_store import MappingStore


@dataclass
class ProcessingResult:
    """Result of document processing."""
    input_path: str
    output_path: str
    total_redactions: int
    redactions_by_type: dict
    uncertain_count: int
    mapping_file: Optional[str] = None


class DocxProcessor:
    """
    Process DOCX files to redact PII.

    Preserves document formatting while replacing PII with placeholders.
    """

    def __init__(
        self,
        detector: PIIDetector,
        mapping_store: MappingStore,
        interactive: bool = True,
        user_callback: Optional[Callable[[PIIMatch], Optional[bool]]] = None,
    ):
        """
        Initialize DOCX processor.

        Args:
            detector: PII detector instance
            mapping_store: Mapping store for tracking replacements
            interactive: Whether to prompt for uncertain detections
            user_callback: Callback for uncertain detections.
                          Returns True (is PII), False (not PII), or None (skip)
        """
        self.detector = detector
        self.mapping_store = mapping_store
        self.interactive = interactive
        self.user_callback = user_callback

    def process(
        self,
        input_path: str,
        output_path: Optional[str] = None,
    ) -> ProcessingResult:
        """
        Process a DOCX file, redacting PII.

        Args:
            input_path: Path to input DOCX file
            output_path: Path for output file. If None, creates <input>_redacted.docx

        Returns:
            ProcessingResult with statistics
        """
        if output_path is None:
            base, ext = os.path.splitext(input_path)
            output_path = f"{base}_redacted{ext}"

        # Load document
        doc = Document(input_path)

        redaction_counts = {}
        total_redactions = 0
        uncertain_count = 0

        # Process all paragraphs
        for paragraph in doc.paragraphs:
            stats = self._process_paragraph(paragraph)
            total_redactions += stats['redacted']
            uncertain_count += stats['uncertain']
            for pii_type, count in stats['by_type'].items():
                redaction_counts[pii_type] = redaction_counts.get(pii_type, 0) + count

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        stats = self._process_paragraph(paragraph)
                        total_redactions += stats['redacted']
                        uncertain_count += stats['uncertain']
                        for pii_type, count in stats['by_type'].items():
                            redaction_counts[pii_type] = redaction_counts.get(pii_type, 0) + count

        # Process headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    stats = self._process_paragraph(paragraph)
                    total_redactions += stats['redacted']
                    uncertain_count += stats['uncertain']
                    for pii_type, count in stats['by_type'].items():
                        redaction_counts[pii_type] = redaction_counts.get(pii_type, 0) + count

            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    stats = self._process_paragraph(paragraph)
                    total_redactions += stats['redacted']
                    uncertain_count += stats['uncertain']
                    for pii_type, count in stats['by_type'].items():
                        redaction_counts[pii_type] = redaction_counts.get(pii_type, 0) + count

        # Save processed document
        doc.save(output_path)

        return ProcessingResult(
            input_path=input_path,
            output_path=output_path,
            total_redactions=total_redactions,
            redactions_by_type=redaction_counts,
            uncertain_count=uncertain_count,
        )

    def _process_paragraph(self, paragraph: Paragraph) -> dict:
        """
        Process a single paragraph, redacting PII.

        Returns stats dict with 'redacted', 'uncertain', 'by_type' keys.
        """
        stats = {'redacted': 0, 'uncertain': 0, 'by_type': {}}

        # Get full text
        full_text = paragraph.text
        if not full_text.strip():
            return stats

        # Detect PII
        result = self.detector.detect(full_text)

        # Collect all matches to process (confirmed + user-approved uncertain)
        matches_to_redact: List[PIIMatch] = list(result.matches)

        # Handle uncertain detections
        for match in result.uncertain:
            if self.interactive and self.user_callback:
                decision = self.user_callback(match)
                if decision is True:
                    matches_to_redact.append(match)
                    # Learn this as PII
                    self.detector.learn_pii(match.text, match.pii_type)
                elif decision is False:
                    # Learn this as safe
                    self.detector.learn_safe(match.text)
                # If None, skip (don't learn)
            stats['uncertain'] += 1

        if not matches_to_redact:
            return stats

        # Sort matches by position (reverse order for replacement)
        matches_to_redact.sort(key=lambda m: m.start, reverse=True)

        # Build replacement map
        replacements: List[Tuple[int, int, str, str]] = []  # (start, end, original, placeholder)

        for match in matches_to_redact:
            placeholder = self.mapping_store.add_mapping(
                match.text,
                match.pii_type.value
            )
            replacements.append((match.start, match.end, match.text, placeholder))

            # Update stats
            stats['redacted'] += 1
            pii_type = match.pii_type.value
            stats['by_type'][pii_type] = stats['by_type'].get(pii_type, 0) + 1

        # Apply replacements to paragraph
        # We need to handle runs carefully to preserve formatting
        self._apply_replacements(paragraph, replacements)

        return stats

    def _apply_replacements(
        self,
        paragraph: Paragraph,
        replacements: List[Tuple[int, int, str, str]]
    ):
        """
        Apply text replacements while preserving run formatting.

        This is complex because DOCX stores text in runs, and a single
        PII match might span multiple runs.
        """
        if not replacements:
            return

        # Get current text and runs
        full_text = paragraph.text

        # Apply all replacements to get new text
        new_text = full_text
        for start, end, original, placeholder in replacements:
            # Recalculate positions as text changes
            offset = len(new_text) - len(full_text)
            adj_start = start + offset
            adj_end = end + offset
            new_text = new_text[:adj_start] + placeholder + new_text[adj_end:]
            full_text = full_text[:start] + placeholder + full_text[end:]

        # Simple approach: preserve first run's formatting, replace all text
        if paragraph.runs:
            # Store first run's formatting
            first_run = paragraph.runs[0]

            # Clear all runs
            for run in paragraph.runs:
                run.text = ""

            # Set new text on first run
            first_run.text = new_text
        else:
            # No runs, just set text directly
            paragraph.text = new_text


class DocxRestorer:
    """Restore redacted DOCX files using mapping store."""

    def __init__(self, mapping_store: MappingStore):
        """
        Initialize restorer.

        Args:
            mapping_store: Mapping store with placeholder->original mappings
        """
        self.mapping_store = mapping_store

    def restore(
        self,
        input_path: str,
        output_path: Optional[str] = None,
    ) -> Tuple[str, int]:
        """
        Restore a redacted DOCX file.

        Args:
            input_path: Path to redacted DOCX file
            output_path: Path for restored file. If None, creates <input>_restored.docx

        Returns:
            Tuple of (output_path, restoration_count)
        """
        if output_path is None:
            base, ext = os.path.splitext(input_path)
            # Remove _redacted suffix if present
            if base.endswith('_redacted'):
                base = base[:-9]
            output_path = f"{base}_restored{ext}"

        # Get all restorations
        restorations = self.mapping_store.get_all_restorations()
        if not restorations:
            raise ValueError("No mappings available for restoration")

        # Load document
        doc = Document(input_path)
        restoration_count = 0

        # Process all paragraphs
        for paragraph in doc.paragraphs:
            count = self._restore_paragraph(paragraph, restorations)
            restoration_count += count

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        count = self._restore_paragraph(paragraph, restorations)
                        restoration_count += count

        # Process headers and footers
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    count = self._restore_paragraph(paragraph, restorations)
                    restoration_count += count

            if section.footer:
                for paragraph in section.footer.paragraphs:
                    count = self._restore_paragraph(paragraph, restorations)
                    restoration_count += count

        # Save restored document
        doc.save(output_path)

        return output_path, restoration_count

    def _restore_paragraph(
        self,
        paragraph: Paragraph,
        restorations: dict
    ) -> int:
        """Restore placeholders in a paragraph."""
        count = 0
        text = paragraph.text

        for placeholder, original in restorations.items():
            if placeholder in text:
                text = text.replace(placeholder, original)
                count += text.count(original)

        if count > 0:
            # Apply restored text
            if paragraph.runs:
                for run in paragraph.runs:
                    run.text = ""
                paragraph.runs[0].text = text
            else:
                paragraph.text = text

        return count
