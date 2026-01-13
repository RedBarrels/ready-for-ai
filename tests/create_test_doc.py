"""Create a test document with various PII for testing."""

from docx import Document

def create_test_docx():
    doc = Document()

    doc.add_heading('Project Phoenix Status Report', 0)

    doc.add_paragraph(
        'From: John Smith (john.smith@acmecorp.com)\n'
        'To: Sarah Johnson (sarah.j@globex.io)\n'
        'CC: @mike_wilson, @jennifer_lee\n'
        'Date: January 10, 2024'
    )

    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This report summarizes the progress of Project Phoenix for Acme Corporation. '
        'Our team lead, Michael Chen, has been coordinating with the client representative '
        'Jennifer Williams at GlobalTech Industries.'
    )

    doc.add_heading('Team Members', level=1)
    doc.add_paragraph(
        '• John Smith - Project Manager (555-123-4567)\n'
        '• Sarah Johnson - Lead Developer\n'
        '• Mike Wilson - Backend Engineer (@mike_wilson)\n'
        '• Jennifer Lee - QA Lead (jennifer.lee@acmecorp.com)\n'
        '• Robert Brown - DevOps (robert.b@acmecorp.com)'
    )

    doc.add_heading('Client Information', level=1)
    doc.add_paragraph(
        'Client: GlobalTech Industries\n'
        'Contact: Jennifer Williams\n'
        'Email: j.williams@globaltech.com\n'
        'Phone: 555-987-6543\n'
        'Address: 123 Main Street, San Francisco, CA 94102'
    )

    doc.add_heading('Technical Details', level=1)
    doc.add_paragraph(
        'The production server is running at IP 192.168.1.100. '
        'API documentation is available at https://api.acmecorp.com/docs. '
        'For support, contact support@acmecorp.com or call 1-800-555-0199.'
    )

    doc.add_heading('Confidential Notes', level=1)
    doc.add_paragraph(
        'Employee SSN for payroll verification: 123-45-6789 (John Smith)\n'
        'Corporate card ending: 4111-1111-1111-1234\n'
        'Internal Slack channel: #project-phoenix-team'
    )

    doc.add_paragraph(
        '\nBest regards,\n'
        'John Smith\n'
        'Project Manager, Engineering Team\n'
        'Acme Corporation'
    )

    doc.save('test_document.docx')
    print('Created test_document.docx')

if __name__ == '__main__':
    create_test_docx()
